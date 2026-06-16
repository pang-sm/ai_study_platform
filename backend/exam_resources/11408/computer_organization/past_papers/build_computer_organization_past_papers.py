"""Build computer_organization past-paper questions from docx with images.
The docx contains question screenshots + answers for 2022-2026.
"""
import json, os, re, sys, io, shutil
from pathlib import Path
from collections import defaultdict

BASE = Path(__file__).resolve().parent.parent.parent.parent.parent
sys.path.insert(0, str(BASE))
from database import SessionLocal
import models

DOCX = BASE / "exam_resources/11408/computer_organization/past_papers/raw/11408_2022_2026_computer_organization_questions_answers.docx"
ASSETS = BASE / "exam_resources/11408/computer_organization/past_papers/assets"
CHKD = BASE / "exam_resources/11408/computer_organization/past_papers/checked"
RPT = BASE / "exam_resources/11408/computer_organization/past_papers/import_reports"

SUBJECT_KEY = "computer_organization"
SUBJECT_NAME = "计算机组成原理"

CHOICE_RANGE = range(12, 23)  # Q12-Q22
BIG_NUMBERS = [43, 44]

def extract_images():
    """Extract all images from docx to assets/{year}/ directories.
    Images are in empty paragraphs right after the heading paragraph.
    """
    from docx import Document

    doc = Document(str(DOCX))
    nsmap_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    nsmap_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    # Build mapping: rel_id -> (year, qnum, img_index)
    img_rel_map = {}
    q_img_counts = defaultdict(int)
    current_year = None
    last_qnum = None

    for p in doc.paragraphs:
        t = p.text.strip()
        ym = re.search(r'(\d{4})\s*年', t) if t else None
        if ym: current_year = int(ym.group(1))

        qm = re.search(r'第\s*(\d+)\s*题', t) if t else None
        if qm: last_qnum = int(qm.group(1))

        # Check for images in this paragraph (text may be empty for image-only paragraphs)
        for blip in p._element.iter(f'{{{nsmap_a}}}blip'):
            embed = blip.get(f'{{{nsmap_r}}}embed')
            if embed and current_year and last_qnum:
                idx = q_img_counts[(current_year, last_qnum)]
                q_img_counts[(current_year, last_qnum)] += 1
                img_rel_map[embed] = (current_year, last_qnum, idx)

    # Actually extract image files
    extracted = defaultdict(list)
    for rel_id, rel in doc.part.rels.items():
        if 'image' not in rel.reltype: continue
        if rel_id not in img_rel_map: continue
        year, qnum, idx = img_rel_map[rel_id]
        ext = rel.target_ref.split('.')[-1]
        if ext.lower() not in ('png', 'jpg', 'jpeg', 'gif', 'bmp'): ext = 'png'
        year_dir = ASSETS / str(year)
        year_dir.mkdir(parents=True, exist_ok=True)
        fname = f'q{qnum:02d}_{idx}.{ext}'
        with open(str(year_dir / fname), 'wb') as f:
            f.write(rel.target_part.blob)
        extracted[(year, qnum)].append(f'assets/{year}/{fname}')

    return extracted, doc, img_rel_map

def parse_questions(extracted_images):
    """Parse questions from docx text + extracted image paths.
    Scans paragraphs sequentially: year → qnum → image → answer.
    """
    from docx import Document
    doc = Document(str(DOCX))

    questions = []
    rejected = []
    current_year = None
    current_qnum = None
    pending_answer = None  # (answer_text, answer_status)

    for p in doc.paragraphs:
        t = p.text.strip()

        # Year header - flush last question of previous year
        if t:
            ym = re.search(r'(\d{4})\s*年', t)
            if ym:
                if current_year and current_qnum and (current_qnum in CHOICE_RANGE or current_qnum in BIG_NUMBERS):
                    qtype = "big" if current_qnum in BIG_NUMBERS else "choice"
                    opts = {} if qtype == "big" else {"A":"(见截图)","B":"(见截图)","C":"(见截图)","D":"(见截图)"}
                    ans, ans_status = pending_answer or ("", "confirmed")
                    imgs = extracted_images.get((current_year, current_qnum), [])
                    questions.append({
                        "exam_type":"11408","subject":SUBJECT_KEY,"course_id":f"{SUBJECT_KEY}_11408",
                        "year":current_year,"paper_name":f"{current_year}年计算机组成原理真题",
                        "question_number":current_qnum,"question_type":qtype,
                        "question_text":f"第{current_qnum}题（见截图）",
                        "options":opts,"answer":ans,"answer_status":ans_status,
                        "explanation":"","image_paths":imgs,
                        "source_ref":f"{current_year}-Q{current_qnum:02d}","is_active":True,
                    })
                current_year = int(ym.group(1))
                current_qnum = None
                pending_answer = None
                continue

            # Question heading - flush previous and start new
            qm = re.search(r'第\s*(\d+)\s*题', t)
            if qm and current_year:
                # Flush previous question
                if current_qnum:
                    qtype = "big" if current_qnum in BIG_NUMBERS else "choice"
                    qtype = qtype if current_qnum in CHOICE_RANGE or current_qnum in BIG_NUMBERS else None
                    if qtype:
                        opts = {} if qtype == "big" else {"A":"(见截图)","B":"(见截图)","C":"(见截图)","D":"(见截图)"}
                        ans, ans_status = pending_answer or ("", "confirmed")
                        imgs = extracted_images.get((current_year, current_qnum), [])
                        questions.append({
                            "exam_type":"11408","subject":SUBJECT_KEY,"course_id":f"{SUBJECT_KEY}_11408",
                            "year":current_year,"paper_name":f"{current_year}年计算机组成原理真题",
                            "question_number":current_qnum,"question_type":qtype,
                            "question_text":f"第{current_qnum}题（见截图）",
                            "options":opts,"answer":ans,"answer_status":ans_status,
                            "explanation":"","image_paths":imgs,
                            "source_ref":f"{current_year}-Q{current_qnum:02d}","is_active":True,
                        })

                current_qnum = int(qm.group(1))
                pending_answer = None
                continue

            # Answer line
            if t.startswith("答案：") and current_qnum:
                ans = t.replace("答案：", "").strip()
                status = "pending" if "待补充" in ans else "confirmed"
                pending_answer = (ans, status)
                continue

    # Flush last question
    if current_year and current_qnum:
        qtype = "big" if current_qnum in BIG_NUMBERS else "choice"
        if qtype in ("choice", "big") and (current_qnum in CHOICE_RANGE or current_qnum in BIG_NUMBERS):
            opts = {} if qtype == "big" else {"A":"(见截图)","B":"(见截图)","C":"(见截图)","D":"(见截图)"}
            ans, ans_status = pending_answer or ("", "confirmed")
            imgs = extracted_images.get((current_year, current_qnum), [])
            questions.append({
                "exam_type":"11408","subject":SUBJECT_KEY,"course_id":f"{SUBJECT_KEY}_11408",
                "year":current_year,"paper_name":f"{current_year}年计算机组成原理真题",
                "question_number":current_qnum,"question_type":qtype,
                "question_text":f"第{current_qnum}题（见截图）",
                "options":opts,"answer":ans,"answer_status":ans_status,
                "explanation":"","image_paths":imgs,
                "source_ref":f"{current_year}-Q{current_qnum:02d}","is_active":True,
            })

    return questions, rejected

def main():
    dry_run = "--dry-run" in sys.argv

    print("Extracting images from docx...")
    extracted_images, doc, img_map = extract_images()
    total_imgs = sum(len(v) for v in extracted_images.values())
    print(f"Extracted {total_imgs} images for {len(extracted_images)} question slots")

    print("Parsing questions...")
    questions, rejected = parse_questions(extracted_images)

    t = len(questions)
    c = sum(1 for q in questions if q["question_type"] == "choice")
    b = t - c
    r = len(rejected)

    print(f"\nParsed: {t} (choice={c}, big={b}, rejected={r})")

    # Per year stats
    years = defaultdict(lambda: {"choice": 0, "big": 0})
    for q in questions:
        years[q["year"]][q["question_type"]] += 1
    for yr in sorted(years):
        y = years[yr]
        print(f"  {yr}: {y['choice']+y['big']} (choice={y['choice']}, big={y['big']})")

    # Validation
    no_answer = [q for q in questions if not q["answer"]]
    pending = [q for q in questions if q["answer_status"] == "pending"]
    no_imgs = [q for q in questions if not q["image_paths"]]
    choice_no_opts = [q for q in questions if q["question_type"] == "choice" and len(q.get("options", {})) == 0]

    if no_answer: print(f"WARNING: {len(no_answer)} questions without answer!")
    if pending: print(f"INFO: {len(pending)} questions with pending answer")
    if no_imgs: print(f"WARNING: {len(no_imgs)} questions without images!")
    if choice_no_opts: print(f"WARNING: {len(choice_no_opts)} choice without options!")

    # Check Q12-Q22 completeness per year
    for yr in sorted(set(q["year"] for q in questions)):
        yr_qs = {q["question_number"] for q in questions if q["year"] == yr}
        missing_choice = [n for n in CHOICE_RANGE if n not in yr_qs]
        missing_big = [n for n in BIG_NUMBERS if n not in yr_qs]
        if missing_choice: print(f"  {yr}: MISSING choice Q{missing_choice}")
        if missing_big: print(f"  {yr}: MISSING big Q{missing_big}")

    # Save
    CHKD.mkdir(parents=True, exist_ok=True)
    RPT.mkdir(parents=True, exist_ok=True)
    json.dump(questions, (CHKD / "parsed_ready.json").open("w", encoding="utf-8"), ensure_ascii=False, indent=2)
    json.dump(rejected, (CHKD / "parsed_rejected.json").open("w", encoding="utf-8"), ensure_ascii=False)

    rp = {
        "total": t, "choice": c, "big": b, "rejected": r,
        "pending_answers": len(pending),
        "no_images": len(no_imgs),
        "per_year": {str(yr): {"total": years[yr]["choice"] + years[yr]["big"], "choice": years[yr]["choice"], "big": years[yr]["big"]} for yr in sorted(years)},
    }
    json.dump(rp, (RPT / "computer_organization_past_papers_report.json").open("w", encoding="utf-8"), ensure_ascii=False, indent=2)
    (RPT / "computer_organization_past_papers_report.md").write_text(
        f"# Past Papers Report\n- total: {t} (choice={c}, big={b})\n- pending answers: {len(pending)}\n" +
        "\n".join(f"- {yr}: {years[yr]['choice']+years[yr]['big']} (c={years[yr]['choice']}, b={years[yr]['big']})" for yr in sorted(years)),
        encoding="utf-8")
    print("Reports saved")

    if dry_run:
        print("DRY-RUN: skipping DB import.")
        return

    # Import to DB
    db = SessionLocal()
    # Deactivate old CO past papers
    db.query(models.ExamQuestionBank).filter(
        models.ExamQuestionBank.subject_key == SUBJECT_KEY,
        models.ExamQuestionBank.source_type == "past_paper"
    ).update({"is_active": False})
    db.commit()

    ins = 0
    for q in questions:
        item = models.ExamQuestionBank(
            subject_key=SUBJECT_KEY, subject_name=SUBJECT_NAME,
            source_type="past_paper", visibility="public",
            knowledge_point_id="", knowledge_point_name="",
            knowledge_point_path="",
            year=q["year"], question_number=q["question_number"],
            question_type=q["question_type"],
            stem=json.dumps({
                "text": q["question_text"],
                "image_paths": q["image_paths"],
                "answer_status": q["answer_status"],
            }, ensure_ascii=False),
            options_json=json.dumps(q.get("options", {}), ensure_ascii=False),
            standard_answer=q.get("answer", ""),
            analysis="",
            difficulty="基础",
            source_ref=f"past_paper:{q['source_ref']}",
            is_active=True,
        )
        db.add(item)
        ins += 1
    db.commit()
    act = db.query(models.ExamQuestionBank).filter(
        models.ExamQuestionBank.subject_key == SUBJECT_KEY,
        models.ExamQuestionBank.source_type == "past_paper",
        models.ExamQuestionBank.is_active == True
    ).count()
    db.close()
    print(f"DB: {ins} inserted, {act} active")

if __name__ == "__main__":
    main()
