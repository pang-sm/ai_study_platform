import sys, os, json, time
sys.path.insert(0, r"C:\Users\26477\Desktop\ai_study_platform\backend")
os.chdir(r"C:\Users\26477\Desktop\ai_study_platform\backend")

from docx import Document
from qwen_parser import parse_image_with_qwen

doc_path = r"exam_resources\11408\data_structure\11408_2022_2026_data_structure_questions_answers.docx"
doc = Document(doc_path)

# Locate 2024 Q1
target_year = "2024"
found = False
for pi, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == target_year + " 年" or text == target_year + "年":
        for pj in range(pi+1, min(pi+60, len(doc.paragraphs))):
            qtext = doc.paragraphs[pj].text.strip()
            if "第 1 题" in qtext or "第1题" in qtext:
                for pk in range(pj+1, min(pj+5, len(doc.paragraphs))):
                    for run in doc.paragraphs[pk].runs:
                        for elem in run._element:
                            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                            if tag in ("drawing", "pict"):
                                for desc in elem.iter():
                                    dtag = desc.tag.split("}")[-1] if "}" in desc.tag else desc.tag
                                    if dtag == "blip":
                                        embed = desc.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                                        if embed and embed in doc.part.rels:
                                            img_blob = doc.part.rels[embed].target_part.blob
                                            img_path = r"cache\feasibility_q1.jpg"
                                            os.makedirs("cache", exist_ok=True)
                                            with open(img_path, "wb") as f:
                                                f.write(img_blob)
                                            print("Image:", img_path, len(img_blob), "bytes")
                                            found = True
                                            
                                            prompt = "识别这张11408数据结构考研真题。提取:1.题号2.题干文字3.A/B/C/D选项。公式代码保留原样。输出JSON。"
                                            t0 = time.time()
                                            result = parse_image_with_qwen(img_path, prompt=prompt, timeout_seconds=60)
                                            print("OCR %.1fs success=%s" % ((time.time()-t0), result.get("success")))
                                            txt = result.get("extracted_text","") or ""
                                            print("Text:", len(txt), "chars")
                                            print("---RAW---")
                                            print(txt[:3000])
                                if found: break
                        if found: break
                if found: break
        if found: break

if not found:
    print("ERROR: 2024 Q1 image not found")
