import sys, os, time
sys.path.insert(0, r"C:\Users\26477\Desktop\ai_study_platform\backend")
os.chdir(r"C:\Users\26477\Desktop\ai_study_platform\backend")
from docx import Document
from qwen_parser import parse_image_with_qwen

doc_path = r"exam_resources\11408\data_structure\11408_2022_2026_data_structure_questions_answers.docx"
doc = Document(doc_path)

# Find 2024 Q41 (big question)
target_year = "2024"
found = False
for pi, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == target_year + " 年" or text == target_year + "年":
        # Find Q41 and its answer
        q41_img = None
        answer = ""
        for pj in range(pi+1, min(pi+60, len(doc.paragraphs))):
            qtext = doc.paragraphs[pj].text.strip()
            if "第 41 题" in qtext or "第41题" in qtext:
                # Extract image
                for pk in range(pj+1, min(pj+5, len(doc.paragraphs))):
                    for run in doc.paragraphs[pk].runs:
                        for elem in run._element:
                            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                            if tag in ("drawing", "pict"):
                                for desc in elem.iter():
                                    dtag = desc.tag.split("}")[-1] if "}" in desc.tag else desc.tag
                                    if dtag == "blip":
                                        embed = desc.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                                        if embed and embed in doc.part.rels:
                                            q41_img = doc.part.rels[embed].target_part.blob
            # Find answer
            if "答案" in qtext and len(qtext) > 3 and not answer:
                ans = qtext.replace("答案：","").replace("答案:","").strip()
                if ans and len(ans) > 5:
                    answer = ans
            
            if q41_img and answer:
                break
        
        if q41_img:
            img_path = r"cache\feasibility_q41.jpg"
            with open(img_path, "wb") as f:
                f.write(q41_img)
            print("Q41 Image:", len(q41_img), "bytes")
            print("Q41 Answer:", answer[:200])
            
            prompt = "识别这张11408数据结构考研大题。提取题干完整文字。保留代码/公式原样。输出JSON: {题干:..., 小题:...}"
            t0 = time.time()
            result = parse_image_with_qwen(img_path, prompt=prompt, timeout_seconds=60)
            print("OCR %.1fs success=%s" % ((time.time()-t0), result.get("success")))
            txt = result.get("extracted_text","") or ""
            print("Text:", len(txt), "chars")
            print("---RAW---")
            print(txt[:3000])
            found = True
        break

if not found:
    print("ERROR: 2024 Q41 not found")
